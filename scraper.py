"""
scraper.py — Web scraping utilities.

Cleaned up in this revision:
  • Removed the dead `build_paraphrased_content` function that tried to
    call `http://127.0.0.1:8000/paraphrase` (an old FastAPI server).
    Paraphrasing is now handled centrally by ai_helper.build_paraphrased_content.
  • Better extraction with main-content heuristics.
  • Saner timeouts and User-Agent headers.
  • Optional max_items cap to avoid huge payloads.
"""

from __future__ import annotations
import re
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt


DEFAULT_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
}


# ─────────────────────────────────────────────
# FETCH HTML
# ─────────────────────────────────────────────

def fetch_html(url: str, timeout: int = 15) -> str:
    """Fetch raw HTML for a URL. Raises on HTTP errors."""
    response = requests.get(url, headers=DEFAULT_HEADERS, timeout=timeout)
    response.raise_for_status()
    response.encoding = response.apparent_encoding or response.encoding
    return response.text


def parse_html(html: str) -> BeautifulSoup:
    return BeautifulSoup(html, "html.parser")


# ─────────────────────────────────────────────
# CLEAN TEXT
# ─────────────────────────────────────────────

def clean_text(text: str) -> str:
    """Normalize whitespace and strip wiki-style citation markers."""
    text = re.sub(r"\[\d+\]", "", text)        # [1] [42]
    text = re.sub(r"\[edit\]", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+", " ", text)           # collapse whitespace
    return text.strip()


# ─────────────────────────────────────────────
# EXTRACT CONTENT
# ─────────────────────────────────────────────

def extract_content(soup: BeautifulSoup, max_items: int = 60) -> list[dict]:
    """
    Extract a list of {type, text} blocks from the main content of a page.
    Skips nav, footer, header, aside, scripts. De-dupes paragraphs.
    """
    container = (
        soup.find("article")
        or soup.find("main")
        or soup.find("div", id=re.compile(r"(content|main|article)", re.I))
        or soup.find("div", class_=lambda x: x and any(
            k in x.lower() for k in ("content", "body", "post", "article", "main")
        ))
        or soup.body
    )

    if not container:
        return []

    # Strip noise
    for tag in container.find_all(["script", "style", "nav", "footer", "header", "aside", "form"]):
        tag.decompose()

    content: list[dict] = []
    allowed_tags = ["h1", "h2", "h3", "h4", "p", "li", "blockquote"]
    seen_texts: set[str] = set()

    for element in container.find_all(allowed_tags):
        if element.find_parent(["nav", "footer", "header", "aside", "script", "style"]):
            continue

        text = clean_text(element.get_text(" ", strip=True))
        if not text or len(text) < 30:
            continue

        key = text.lower()
        if key in seen_texts:
            continue
        seen_texts.add(key)

        content.append({"type": element.name, "text": text})

        if len(content) >= max_items:
            break

    return content


# ─────────────────────────────────────────────
# WORD DOCUMENT EXPORT
# ─────────────────────────────────────────────

def _add_block(document: Document, tag_type: str, text: str):
    """Add a single content block to the docx based on its tag type."""
    if not text:
        return

    if tag_type.startswith("h") and len(tag_type) == 2 and tag_type[1].isdigit():
        level = max(1, min(int(tag_type[1]) - 1, 4))
        document.add_heading(text, level=level)
    elif tag_type == "li":
        document.add_paragraph(text, style="List Bullet")
    elif tag_type in ("p", "blockquote"):
        para = document.add_paragraph(text)
        para.paragraph_format.space_after = Pt(8)
    else:
        document.add_paragraph(text)


def save_to_word(content: list[dict], filename: str = "scraped_content.docx"):
    """Save a list of {type, text} blocks to a Word document."""
    document = Document()
    for item in content:
        _add_block(document, item.get("type", "p"), item.get("text", ""))
    document.save(filename)


def save_paraphrased_to_word(content: list[dict], filename: str = "paraphrased_content.docx"):
    """Save the paraphrased version of each block to Word."""
    document = Document()
    for item in content:
        text = item.get("paraphrased") or item.get("text", "")
        _add_block(document, item.get("type", "p"), text)
    document.save(filename)

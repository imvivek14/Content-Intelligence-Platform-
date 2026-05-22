"""
AI Chat — ChatGPT/Claude-style conversation interface.

Features:
  • Streaming responses (tokens appear as they're generated)
  • Persistent conversation history within the session
  • Custom system prompt (set the AI's personality/role)
  • Quick-start preset personas (tutor, coder, writer, etc.)
  • New Chat / Clear / Export-as-DOCX buttons
  • Visible model indicator
  • Real error display (no silent failures)
"""

from __future__ import annotations
import io
import streamlit as st
from docx import Document

from ai_helper import chat_stream, get_active_model, render_api_key_sidebar, _resolve_api_key
from styles import inject_css, page_header, gradient_divider


# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────

st.set_page_config(page_title="AI Chat", layout="wide", page_icon="💬")
inject_css()
render_api_key_sidebar()

page_header(
    icon="💬",
    title="AI Chat",
    subtitle="Have a real conversation with Gemini — ChatGPT-style streaming, conversation memory, and custom personas.",
    badge="Module 00 · Conversational AI",
)


# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────

if "chat_messages" not in st.session_state:
    st.session_state["chat_messages"] = []     # list of {role, content}

if "chat_system_prompt" not in st.session_state:
    st.session_state["chat_system_prompt"] = ""


# ─────────────────────────────────────────────
# SIDEBAR EXTRAS — SYSTEM PROMPT & ACTIONS
# ─────────────────────────────────────────────

with st.sidebar:
    st.markdown("### 💬 Chat Settings")

    # Persona presets
    persona = st.selectbox(
        "Quick persona",
        options=[
            "🤖 Default assistant",
            "👨‍🏫 Patient tutor",
            "💻 Senior software engineer",
            "✍️ Creative writer",
            "🔬 Research analyst",
            "🎯 Career coach",
            "🍳 Friendly cook",
            "Custom (write your own)",
        ],
        key="chat_persona_select",
    )

    persona_prompts = {
        "🤖 Default assistant":
            "You are a helpful, harmless, and honest AI assistant. Be concise and clear.",
        "👨‍🏫 Patient tutor":
            "You are a patient tutor. Explain things step-by-step using simple language and examples. "
            "Check the user's understanding before moving on.",
        "💻 Senior software engineer":
            "You are a senior software engineer. Give precise, idiomatic code with brief explanations. "
            "Always point out edge cases and suggest tests.",
        "✍️ Creative writer":
            "You are a creative writer with a vivid imagination. Use rich imagery, varied sentence "
            "lengths, and strong verbs.",
        "🔬 Research analyst":
            "You are a careful research analyst. Cite reasoning, distinguish facts from opinions, "
            "and acknowledge uncertainty when relevant.",
        "🎯 Career coach":
            "You are an encouraging career coach. Ask clarifying questions, listen actively, and "
            "give concrete, actionable advice.",
        "🍳 Friendly cook":
            "You are a friendly home cook. Suggest practical recipes with common ingredients and "
            "explain techniques clearly.",
    }

    if persona == "Custom (write your own)":
        custom_prompt = st.text_area(
            "Custom system prompt",
            value=st.session_state.get("chat_custom_prompt", ""),
            height=140,
            placeholder="e.g. 'You are a Shakespearean poet who only speaks in iambic pentameter...'",
            key="chat_custom_prompt",
        )
        st.session_state["chat_system_prompt"] = custom_prompt
    else:
        st.session_state["chat_system_prompt"] = persona_prompts.get(persona, "")
        with st.expander("Preview persona prompt"):
            st.write(st.session_state["chat_system_prompt"])

    st.markdown("---")

    # Action buttons
    cb1, cb2 = st.columns(2)
    with cb1:
        if st.button("🆕 New Chat", use_container_width=True, key="chat_new"):
            st.session_state["chat_messages"] = []
            st.rerun()
    with cb2:
        if st.button("🗑️ Clear", use_container_width=True, key="chat_clear",
                     disabled=not st.session_state["chat_messages"]):
            st.session_state["chat_messages"] = []
            st.rerun()


# ─────────────────────────────────────────────
# TOP BAR — MODEL & MESSAGE COUNT
# ─────────────────────────────────────────────

has_key = bool(_resolve_api_key())
working_model = st.session_state.get("_gemini_working_model")

c1, c2, c3 = st.columns([1, 1, 2])
with c1:
    st.markdown(
        f'<span class="status-pill success">🤖 {working_model or get_active_model()}</span>',
        unsafe_allow_html=True,
    )
with c2:
    msg_count = len(st.session_state["chat_messages"])
    st.markdown(
        f'<span class="status-pill success">💬 {msg_count} messages</span>',
        unsafe_allow_html=True,
    )

if not has_key:
    st.error("⚠️ Add your Gemini API key in the sidebar to start chatting. Then click **🔌 Test Connection** to verify it works.")


# ─────────────────────────────────────────────
# QUICK-START SUGGESTIONS (only shown on empty chat)
# ─────────────────────────────────────────────

if not st.session_state["chat_messages"] and has_key:
    gradient_divider()
    st.markdown("##### 💡 Try one of these to start")

    suggestions = [
        "Explain quantum entanglement in simple terms",
        "Write a 3-day Tokyo travel itinerary",
        "Help me debug a Python recursion error",
        "Suggest a healthy weekly meal plan",
        "Brainstorm names for a coffee shop",
        "Summarize the plot of Hamlet in 5 sentences",
    ]

    sug_cols = st.columns(3)
    for i, s in enumerate(suggestions):
        with sug_cols[i % 3]:
            if st.button(s, use_container_width=True, key=f"sug_{i}"):
                st.session_state["chat_messages"].append({"role": "user", "content": s})
                st.session_state["_pending_response"] = True
                st.rerun()


# ─────────────────────────────────────────────
# MESSAGE HISTORY DISPLAY
# ─────────────────────────────────────────────

gradient_divider()
chat_container = st.container()
with chat_container:
    for msg in st.session_state["chat_messages"]:
        with st.chat_message(msg["role"], avatar="🧑" if msg["role"] == "user" else "✨"):
            st.markdown(msg["content"])


# ─────────────────────────────────────────────
# STREAMING RESPONSE FOR PENDING USER MESSAGE
# ─────────────────────────────────────────────

if st.session_state.get("_pending_response") and st.session_state["chat_messages"]:
    last = st.session_state["chat_messages"][-1]
    if last["role"] == "user":
        # Build messages list with system prompt
        messages_for_api: list[dict] = []
        sys_prompt = st.session_state.get("chat_system_prompt", "").strip()
        if sys_prompt:
            messages_for_api.append({"role": "system", "content": sys_prompt})
        messages_for_api.extend(st.session_state["chat_messages"])

        # Stream the assistant response
        with st.chat_message("assistant", avatar="✨"):
            placeholder = st.empty()
            full_response = ""
            try:
                for chunk in chat_stream(messages_for_api):
                    full_response += chunk
                    # Show with a blinking caret while streaming
                    placeholder.markdown(full_response + "▌")
                placeholder.markdown(full_response if full_response else "_(empty response)_")
            except Exception as e:
                placeholder.error(f"❌ Error: {e}")
                full_response = f"⚠️ Error: {e}"

        # Save to history
        st.session_state["chat_messages"].append({
            "role": "assistant",
            "content": full_response or "_(empty response)_",
        })
        st.session_state.pop("_pending_response", None)
        st.rerun()


# ─────────────────────────────────────────────
# CHAT INPUT
# ─────────────────────────────────────────────

user_input = st.chat_input("Type your message...", disabled=not has_key)

if user_input:
    st.session_state["chat_messages"].append({"role": "user", "content": user_input})
    st.session_state["_pending_response"] = True
    st.rerun()


# ─────────────────────────────────────────────
# EXPORT
# ─────────────────────────────────────────────

if st.session_state["chat_messages"]:
    gradient_divider()

    def export_chat() -> bytes:
        doc = Document()
        doc.add_heading("AI Chat Conversation", level=1)
        sys_prompt = st.session_state.get("chat_system_prompt", "").strip()
        if sys_prompt:
            doc.add_heading("System Prompt", level=2)
            doc.add_paragraph(sys_prompt)
            doc.add_heading("Conversation", level=2)
        for msg in st.session_state["chat_messages"]:
            role = "User" if msg["role"] == "user" else "Assistant"
            p = doc.add_paragraph()
            p.add_run(f"{role}: ").bold = True
            p.add_run(msg["content"])
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    st.download_button(
        "📥 Export Conversation (.docx)",
        data=export_chat(),
        file_name="ai_chat_conversation.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

st.caption("AI Chat · Streaming responses · Powered by Google Gemini")

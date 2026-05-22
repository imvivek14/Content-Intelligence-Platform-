"""
URL Scraper — Extract content from any URL, AI-paraphrase, summarize, and download.

Fully automated: one click runs the entire pipeline:
  scrape → paraphrase → summarize → extract keywords → ready to download.
"""

import io
import streamlit as st
from docx import Document
from docx.shared import Pt

from scraper import fetch_html, parse_html, extract_content
from ai_helper import (
    paraphrase_text,
    summarize_text,
    extract_keywords,
    render_api_key_sidebar,
)
from styles import inject_css, page_header, gradient_divider


# ─────────────────────────────────────────────
# PAGE
# ─────────────────────────────────────────────

st.set_page_config(page_title="URL Scraper", layout="wide", page_icon="🌐")
inject_css()
render_api_key_sidebar()

page_header(
    icon="🌐",
    title="AI URL Scraper",
    subtitle="Paste a URL → automatic scrape, paraphrase, summary, and keyword extraction. Download everything as Word docs.",
    badge="Module 01",
)


# ─────────────────────────────────────────────
# INPUT
# ─────────────────────────────────────────────

with st.container():
    url = st.text_input(
        "🔗 Website URL",
        placeholder="https://en.wikipedia.org/wiki/Artificial_intelligence",
        key="url_input",
    )

    c1, c2, c3 = st.columns([2, 1, 1])
    with c1:
        max_sections = st.slider(
            "Sections to AI-paraphrase",
            min_value=3, max_value=20, value=8,
            help="Higher values use more API quota but give richer output.",
        )
    with c2:
        auto_run = st.checkbox("⚡ Auto-run full pipeline", value=True, help="Scrape + paraphrase + summarize + keywords in one click.")
    with c3:
        st.write("")
        st.write("")
        run_btn = st.button("🚀 Run", use_container_width=True, type="primary")

st.caption("💡 Works best with Wikipedia, news articles, and blog posts.")
gradient_divider()


# ─────────────────────────────────────────────
# AUTOMATED PIPELINE
# ─────────────────────────────────────────────

def run_full_pipeline(url: str, max_sections: int):
    """Scrape → paraphrase → summarize → keywords. Stores results in session_state."""

    progress = st.progress(0, text="Starting pipeline...")

    # 1. Scrape
    progress.progress(10, text="🌐 Fetching page HTML...")
    html = fetch_html(url)

    progress.progress(25, text="📄 Parsing and extracting content...")
    soup = parse_html(html)
    content = extract_content(soup)

    if not content:
        progress.empty()
        st.warning("⚠️ No content was extracted. The page may be dynamic, paywalled, or blocking scrapers.")
        return

    st.session_state["scraped_content"] = content
    st.session_state.pop("paraphrased_sections", None)
    st.session_state.pop("ai_summary", None)
    st.session_state.pop("ai_keywords", None)

    progress.progress(40, text=f"✅ Extracted {len(content)} sections — starting AI...")

    # If no API key, stop here
    if not st.session_state.get("gemini_api_key", "").strip():
        progress.empty()
        st.warning("⚠️ No Gemini API key — extracted content only. Add a key in the sidebar to enable AI steps.")
        return

    # 2. Paraphrase top sections
    paraphrased = []
    sections_to_paraphrase = content[:max_sections]
    for i, item in enumerate(sections_to_paraphrase):
        original = item.get("text", "")
        pct = 40 + int((i / max(len(sections_to_paraphrase), 1)) * 30)
        progress.progress(pct, text=f"🤖 AI paraphrasing section {i + 1}/{len(sections_to_paraphrase)}...")
        if len(original.strip()) > 30:
            new_text = paraphrase_text(original)
        else:
            new_text = original
        paraphrased.append({**item, "paraphrased": new_text})
    st.session_state["paraphrased_sections"] = paraphrased

    # 3. Summarize full page
    progress.progress(75, text="📝 Generating AI summary...")
    full_text = "\n\n".join(item.get("text", "") for item in content if len(item.get("text", "")) > 30)
    st.session_state["ai_summary"] = summarize_text(full_text, style="detailed")

    # 4. Keywords
    progress.progress(90, text="🔑 Extracting keywords...")
    st.session_state["ai_keywords"] = extract_keywords(full_text, n=10)

    progress.progress(100, text="✅ Pipeline complete!")
    progress.empty()


if run_btn:
    if not url.strip():
        st.error("⚠️ Please enter a URL.")
    elif not url.startswith(("http://", "https://")):
        st.error("⚠️ URL must start with http:// or https://")
    else:
        try:
            if auto_run:
                run_full_pipeline(url, max_sections)
                st.success("🎉 Full pipeline complete — scroll down to see results.")
            else:
                # Manual mode — just scrape
                with st.spinner("Scraping..."):
                    html = fetch_html(url)
                    soup = parse_html(html)
                    content = extract_content(soup)
                    st.session_state["scraped_content"] = content
                    st.session_state.pop("paraphrased_sections", None)
                    st.session_state.pop("ai_summary", None)
                    st.session_state.pop("ai_keywords", None)
                if content:
                    st.success(f"✅ Extracted {len(content)} sections (manual mode — click buttons below for AI).")
                else:
                    st.warning("⚠️ No content extracted.")
        except Exception as e:
            st.error(f"❌ Error: {e}")


# ─────────────────────────────────────────────
# MANUAL CONTROL BUTTONS (for non-auto mode)
# ─────────────────────────────────────────────

if "scraped_content" in st.session_state and not auto_run:
    has_key = bool(st.session_state.get("gemini_api_key", "").strip())
    b1, b2, b3 = st.columns(3)
    if b1.button("🤖 Paraphrase All", use_container_width=True, disabled=not has_key):
        with st.spinner("Paraphrasing..."):
            paraphrased = []
            for item in st.session_state["scraped_content"][:max_sections]:
                pt = paraphrase_text(item.get("text", "")) if len(item.get("text", "")) > 30 else item.get("text", "")
                paraphrased.append({**item, "paraphrased": pt})
            st.session_state["paraphrased_sections"] = paraphrased
        st.rerun()
    if b2.button("📝 AI Summarize", use_container_width=True, disabled=not has_key):
        with st.spinner("Summarizing..."):
            full_text = "\n\n".join(it.get("text", "") for it in st.session_state["scraped_content"] if len(it.get("text", "")) > 30)
            st.session_state["ai_summary"] = summarize_text(full_text, style="detailed")
        st.rerun()
    if b3.button("🔑 Extract Keywords", use_container_width=True, disabled=not has_key):
        with st.spinner("Extracting keywords..."):
            full_text = "\n\n".join(it.get("text", "") for it in st.session_state["scraped_content"] if len(it.get("text", "")) > 30)
            st.session_state["ai_keywords"] = extract_keywords(full_text, n=10)
        st.rerun()


# ─────────────────────────────────────────────
# RESULTS — METRICS
# ─────────────────────────────────────────────

if "scraped_content" in st.session_state:
    content = st.session_state["scraped_content"]
    full_text = " ".join(item.get("text", "") for item in content)
    word_count = len(full_text.split())

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("📦 Sections", len(content))
    m2.metric("📝 Words", f"{word_count:,}")
    m3.metric("🤖 Paraphrased", len(st.session_state.get("paraphrased_sections", [])))
    m4.metric("🔑 Keywords", len(st.session_state.get("ai_keywords", [])))

    gradient_divider()


# ─────────────────────────────────────────────
# RESULTS — KEYWORDS
# ─────────────────────────────────────────────

if "ai_keywords" in st.session_state and st.session_state["ai_keywords"]:
    st.markdown("### 🔑 Extracted Keywords")
    kw_html = " ".join(
        f'<span class="status-pill success" style="font-size:0.85rem; margin:4px;">{kw}</span>'
        for kw in st.session_state["ai_keywords"]
    )
    st.markdown(f'<div style="line-height:2.4;">{kw_html}</div>', unsafe_allow_html=True)
    st.write("")


# ─────────────────────────────────────────────
# RESULTS — SUMMARY
# ─────────────────────────────────────────────

if "ai_summary" in st.session_state and st.session_state["ai_summary"]:
    st.markdown("### 📝 AI Page Summary")
    st.markdown(
        f'<div class="content-box ai">{st.session_state["ai_summary"]}</div>',
        unsafe_allow_html=True,
    )
    st.write("")


# ─────────────────────────────────────────────
# RESULTS — SIDE BY SIDE PREVIEW
# ─────────────────────────────────────────────

if "scraped_content" in st.session_state and "paraphrased_sections" in st.session_state:
    st.markdown("### 📊 Original vs AI Paraphrased")

    tab1, tab2 = st.tabs(["📄 Side-by-side", "🔍 Full content"])

    with tab1:
        col_l, col_r = st.columns(2, gap="medium")
        with col_l:
            st.markdown("#### 📄 Original")
            for item in st.session_state["scraped_content"][:6]:
                text = item.get("text", "")
                if text:
                    st.markdown(f'<div class="content-box original">{text}</div>', unsafe_allow_html=True)
        with col_r:
            st.markdown("#### 🤖 AI Paraphrased")
            for item in st.session_state["paraphrased_sections"][:6]:
                text = item.get("paraphrased", item.get("text", ""))
                if text and len(text) > 20:
                    st.markdown(f'<div class="content-box ai">{text}</div>', unsafe_allow_html=True)

    with tab2:
        for i, item in enumerate(st.session_state["scraped_content"]):
            text = item.get("text", "")
            if text:
                tag = item.get("type", "p")
                st.markdown(f"**[{tag}]** {text}")


# ─────────────────────────────────────────────
# DOWNLOAD
# ─────────────────────────────────────────────

def make_doc_bytes(sections, key="text", title: str = ""):
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for item in sections:
        text = item.get(key) or item.get("text", "")
        t = item.get("type", "p")
        if not text or not text.strip():
            continue
        if t.startswith("h"):
            doc.add_heading(text, level=2)
        elif t == "li":
            doc.add_paragraph(text, style="List Bullet")
        else:
            para = doc.add_paragraph(text)
            para.paragraph_format.space_after = Pt(8)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def make_summary_doc(summary: str, keywords: list, title: str = "AI Summary"):
    doc = Document()
    doc.add_heading(title, level=1)
    if keywords:
        doc.add_heading("Keywords", level=2)
        doc.add_paragraph(", ".join(keywords))
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(summary)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


if "scraped_content" in st.session_state:
    gradient_divider()
    st.markdown("### 📥 Download")

    cols = st.columns(3)
    with cols[0]:
        st.download_button(
            "📄 Original",
            make_doc_bytes(st.session_state["scraped_content"], "text", title="Scraped Content"),
            file_name="scraped_original.docx",
            use_container_width=True,
        )
    if "paraphrased_sections" in st.session_state:
        with cols[1]:
            st.download_button(
                "🤖 AI Paraphrased",
                make_doc_bytes(st.session_state["paraphrased_sections"], "paraphrased", title="AI Paraphrased Content"),
                file_name="ai_paraphrased.docx",
                use_container_width=True,
            )
    if "ai_summary" in st.session_state:
        with cols[2]:
            st.download_button(
                "📝 Summary + Keywords",
                make_summary_doc(
                    st.session_state["ai_summary"],
                    st.session_state.get("ai_keywords", []),
                ),
                file_name="ai_summary.docx",
                use_container_width=True,
            )

gradient_divider()
st.caption("URL Scraper · Powered by Google Gemini 2.0 Flash · BeautifulSoup4")

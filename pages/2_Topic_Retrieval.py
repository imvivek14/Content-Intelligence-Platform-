"""
Topic Retrieval — Search any topic across multiple sources, get an AI-written overview.

Pulls from Wikipedia, Simple Wikipedia, Britannica, and Scholarpedia, then asks
Gemini to synthesize a clean structured overview from all of them.
"""

import io
import random
import re
import streamlit as st
from docx import Document

from scraper import fetch_html, parse_html, extract_content
from ai_helper import (
    ai_summarize_topic,
    paraphrase_text,
    extract_keywords,
    summarize_text,
    render_api_key_sidebar,
)
from styles import inject_css, page_header, gradient_divider


# ─────────────────────────────────────────────
# PAGE
# ─────────────────────────────────────────────

st.set_page_config(page_title="Topic Retrieval", layout="wide", page_icon="🔍")
inject_css()
render_api_key_sidebar()

page_header(
    icon="🔍",
    title="AI Topic Retrieval",
    subtitle="Pull content from multiple knowledge sources at once. Gemini AI synthesizes it into a clean, structured overview.",
    badge="Module 02",
)


# ─────────────────────────────────────────────
# SOURCE REGISTRY
# ─────────────────────────────────────────────

def build_sources(topic: str) -> list[dict]:
    t_wiki = topic.replace(" ", "_")
    t_query = topic.replace(" ", "+")
    return [
        {"name": "Wikipedia",        "url": f"https://en.wikipedia.org/wiki/{t_wiki}"},
        {"name": "Simple Wikipedia", "url": f"https://simple.wikipedia.org/wiki/{t_wiki}"},
        {"name": "Britannica",       "url": f"https://www.britannica.com/search?query={t_query}"},
        {"name": "Scholarpedia",     "url": f"http://www.scholarpedia.org/article/{t_wiki}"},
    ]


def clean_and_filter(content: list[dict]) -> list[dict]:
    """Drop boilerplate, navigation, very short blocks, and duplicates."""
    filtered = []
    seen = set()
    boilerplate_patterns = [
        r"cookie", r"privacy policy", r"terms of (use|service)",
        r"all rights reserved", r"subscribe", r"newsletter",
        r"advertisement", r"skip to", r"jump to", r"edit this",
        r"^\s*\[",
    ]
    for item in content:
        text = item.get("text", "")
        if len(text) < 40:
            continue
        lower = text.lower()
        if any(re.search(p, lower) for p in boilerplate_patterns):
            continue
        if lower in seen:
            continue
        seen.add(lower)
        filtered.append(item)
    return filtered


def fetch_raw_text(topic: str) -> tuple[str, list[str]]:
    """Try each source in order; collect text until we have enough."""
    sources = build_sources(topic)
    all_text_parts: list[str] = []
    sources_used: list[str] = []

    for source in sources:
        try:
            html = fetch_html(source["url"], timeout=12)
            soup = parse_html(html)
            extracted = extract_content(soup)
            filtered = clean_and_filter(extracted)
            if len(filtered) >= 3:
                sample = random.sample(filtered, min(8, len(filtered)))
                all_text_parts.extend([item["text"] for item in sample])
                sources_used.append(source["name"])
        except Exception:
            continue
        if len(all_text_parts) >= 18:
            break

    return "\n\n".join(all_text_parts), sources_used


# ─────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────

topic = st.text_input(
    "🔎 Topic",
    placeholder="e.g. Artificial Intelligence, Climate Change, Black Holes",
)

c1, c2, c3 = st.columns([2, 1, 1])
with c1:
    auto_pipeline = st.checkbox("⚡ Auto-run full pipeline", value=True, help="Fetch + summarize + paraphrase + keywords in one click.")
with c2:
    st.write("")
    st.write("")
    fetch_btn = st.button("🚀 Run", use_container_width=True, type="primary")

gradient_divider()


# ─────────────────────────────────────────────
# AUTOMATED PIPELINE
# ─────────────────────────────────────────────

def run_topic_pipeline(topic: str):
    progress = st.progress(0, text="Starting...")

    progress.progress(15, text=f"🌐 Fetching content for '{topic}' from web sources...")
    raw_text, sources_used = fetch_raw_text(topic)

    if not raw_text.strip():
        progress.empty()
        st.error("❌ Could not retrieve content. Try a more specific or commonly-known topic.")
        return

    st.session_state["topic_raw_text"] = raw_text
    st.session_state["topic_sources_used"] = sources_used
    st.session_state.pop("topic_overview", None)
    st.session_state.pop("topic_paraphrased", None)
    st.session_state.pop("topic_keywords", None)
    st.session_state.pop("topic_tldr", None)

    if not st.session_state.get("gemini_api_key", "").strip():
        progress.empty()
        st.warning("⚠️ Raw content fetched, but no Gemini API key set. Add one to enable AI overview.")
        return

    progress.progress(45, text="🤖 Gemini is writing a structured overview...")
    overview = ai_summarize_topic(topic, raw_text)
    st.session_state["topic_overview"] = overview

    progress.progress(70, text="📝 Generating TL;DR...")
    st.session_state["topic_tldr"] = summarize_text(overview, style="tldr")

    progress.progress(85, text="🔑 Extracting keywords...")
    st.session_state["topic_keywords"] = extract_keywords(overview, n=8)

    progress.progress(100, text="✅ Done!")
    progress.empty()


if fetch_btn:
    if not topic.strip():
        st.error("⚠️ Please enter a topic.")
    else:
        try:
            if auto_pipeline:
                run_topic_pipeline(topic.strip())
                if "topic_overview" in st.session_state:
                    st.success(f"🎉 Done! Sources: {', '.join(st.session_state.get('topic_sources_used', []))}")
            else:
                # Manual: just fetch raw + write overview
                with st.spinner(f"Fetching '{topic}'..."):
                    raw_text, sources_used = fetch_raw_text(topic.strip())
                if not raw_text.strip():
                    st.error("❌ Could not retrieve content.")
                else:
                    st.session_state["topic_raw_text"] = raw_text
                    st.session_state["topic_sources_used"] = sources_used
                    if st.session_state.get("gemini_api_key", "").strip():
                        with st.spinner("🤖 Generating AI overview..."):
                            st.session_state["topic_overview"] = ai_summarize_topic(topic.strip(), raw_text)
                        st.success(f"✅ Sources: {', '.join(sources_used)}")
                    else:
                        st.warning("⚠️ Raw content fetched. Add API key to generate AI overview.")
        except Exception as e:
            st.error(f"❌ Error: {e}")


# ─────────────────────────────────────────────
# RESULTS
# ─────────────────────────────────────────────

if "topic_overview" in st.session_state:
    sources = st.session_state.get("topic_sources_used", [])
    if sources:
        sources_html = " ".join(
            f'<span class="status-pill success" style="margin-right:6px;">📚 {s}</span>'
            for s in sources
        )
        st.markdown(f"**Sources used:** {sources_html}", unsafe_allow_html=True)
        st.write("")

    # TL;DR card
    if "topic_tldr" in st.session_state and st.session_state["topic_tldr"]:
        st.markdown(
            f"""
            <div class="glass-card" style="border-left: 4px solid #fbbf24;">
                <div style="color:#fbbf24; font-weight:700; font-size:0.85rem; text-transform:uppercase; letter-spacing:0.06em; margin-bottom:0.4rem;">
                    💡 TL;DR
                </div>
                <div style="color:#e2e8f0; font-size:1.05rem; line-height:1.6;">
                    {st.session_state["topic_tldr"]}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.write("")

    # Keywords
    if "topic_keywords" in st.session_state and st.session_state["topic_keywords"]:
        st.markdown("#### 🔑 Key Concepts")
        kw_html = " ".join(
            f'<span class="status-pill success" style="font-size:0.85rem; margin:4px;">{kw}</span>'
            for kw in st.session_state["topic_keywords"]
        )
        st.markdown(f'<div style="line-height:2.4;">{kw_html}</div>', unsafe_allow_html=True)
        st.write("")

    # Overview
    st.markdown(f"### 🤖 AI Overview: {topic}")
    st.markdown(st.session_state["topic_overview"])

    # Paraphrase button
    if st.button("✍️ AI Paraphrase This Overview", use_container_width=False):
        with st.spinner("Paraphrasing..."):
            st.session_state["topic_paraphrased"] = paraphrase_text(st.session_state["topic_overview"])

    if "topic_paraphrased" in st.session_state:
        gradient_divider()
        st.markdown("### ✍️ Paraphrased Version")
        st.markdown(st.session_state["topic_paraphrased"])

    with st.expander("📄 View Raw Scraped Content"):
        raw = st.session_state.get("topic_raw_text", "")
        st.text(raw[:3000] + ("..." if len(raw) > 3000 else ""))


# ─────────────────────────────────────────────
# DOWNLOAD
# ─────────────────────────────────────────────

def make_doc_bytes(text: str, title: str = "", keywords: list | None = None):
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    if keywords:
        doc.add_paragraph(f"Keywords: {', '.join(keywords)}", style="Intense Quote")

    for line in (text or "").split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


if "topic_overview" in st.session_state:
    gradient_divider()
    st.markdown("### 📥 Download")

    dl_cols = st.columns(2)
    with dl_cols[0]:
        st.download_button(
            "📄 Download Overview",
            make_doc_bytes(
                st.session_state["topic_overview"],
                title=f"AI Overview: {topic}",
                keywords=st.session_state.get("topic_keywords"),
            ),
            file_name=f"topic_{topic.replace(' ', '_')}_overview.docx",
            use_container_width=True,
        )
    if "topic_paraphrased" in st.session_state:
        with dl_cols[1]:
            st.download_button(
                "✍️ Download Paraphrased",
                make_doc_bytes(st.session_state["topic_paraphrased"], title=f"Paraphrased: {topic}"),
                file_name=f"topic_{topic.replace(' ', '_')}_paraphrased.docx",
                use_container_width=True,
            )

gradient_divider()
st.caption("Topic Retrieval · Multi-source web fetch · Google Gemini 2.0 Flash synthesis")

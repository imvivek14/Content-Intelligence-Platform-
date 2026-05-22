"""
Smart Pipeline — Fully automated all-in-one workflow.

Accepts ANY input type (URL, topic, raw text, or prompt) and runs the
full AI pipeline automatically: extract → analyze → paraphrase → summarize
→ keywords → enhance → ready to download.
"""

import io
import streamlit as st
from docx import Document

from scraper import fetch_html, parse_html, extract_content
from ai_helper import (
    paraphrase_text,
    summarize_text,
    enhance_text,
    generate_content,
    extract_keywords,
    detect_language_and_topic,
    ai_summarize_topic,
    render_api_key_sidebar,
)
from styles import inject_css, page_header, gradient_divider


# ─────────────────────────────────────────────
# PAGE
# ─────────────────────────────────────────────

st.set_page_config(page_title="Smart Pipeline", layout="wide", page_icon="🎯")
inject_css()
render_api_key_sidebar()

page_header(
    icon="🎯",
    title="Smart Pipeline",
    subtitle="One workflow, four input types. Drop in a URL, topic, raw text, or prompt — Gemini handles the rest end-to-end.",
    badge="Full Automation",
)


# ─────────────────────────────────────────────
# INPUT TYPE SELECTOR
# ─────────────────────────────────────────────

input_type = st.radio(
    "Choose input type",
    options=["🌐 URL", "🔍 Topic", "📝 Raw Text", "💬 Prompt"],
    horizontal=True,
    key="sp_input_type",
)

placeholder_map = {
    "🌐 URL":      "https://en.wikipedia.org/wiki/Artificial_intelligence",
    "🔍 Topic":    "e.g. Quantum Computing, Black Holes, Renewable Energy",
    "📝 Raw Text": "Paste any text here — an essay, an email, an article, anything.",
    "💬 Prompt":   "e.g. 'Write a 500-word essay on the future of remote work'",
}

if input_type == "📝 Raw Text":
    user_input = st.text_area("Input", height=180, placeholder=placeholder_map[input_type], key="sp_input_text")
else:
    user_input = st.text_input("Input", placeholder=placeholder_map[input_type], key="sp_input_short")


run_btn = st.button("🚀 Run Smart Pipeline", use_container_width=True, type="primary")


# ─────────────────────────────────────────────
# PIPELINE EXECUTION
# ─────────────────────────────────────────────

def get_text_from_input(input_type: str, value: str, progress) -> str | None:
    """Step 1: Convert any input type to a body of text we can analyze."""

    if input_type == "🌐 URL":
        if not value.startswith(("http://", "https://")):
            st.error("⚠️ URL must start with http:// or https://")
            return None
        progress.progress(10, text="🌐 Fetching URL...")
        html = fetch_html(value)
        progress.progress(20, text="📄 Extracting content...")
        soup = parse_html(html)
        content = extract_content(soup)
        if not content:
            st.error("⚠️ Could not extract content from this URL.")
            return None
        return "\n\n".join(item.get("text", "") for item in content)

    if input_type == "🔍 Topic":
        # Simple Wikipedia fetch
        progress.progress(10, text=f"🔍 Fetching content on '{value}'...")
        try:
            html = fetch_html(f"https://en.wikipedia.org/wiki/{value.replace(' ', '_')}")
        except Exception:
            html = fetch_html(f"https://simple.wikipedia.org/wiki/{value.replace(' ', '_')}")
        progress.progress(20, text="📄 Extracting...")
        soup = parse_html(html)
        content = extract_content(soup)
        if not content:
            st.error("⚠️ Could not retrieve information about this topic.")
            return None
        return "\n\n".join(item.get("text", "") for item in content[:25])

    if input_type == "📝 Raw Text":
        progress.progress(20, text="📝 Reading text...")
        return value

    if input_type == "💬 Prompt":
        progress.progress(15, text="🤖 Generating content from prompt...")
        return generate_content(value)

    return None


def run_pipeline(input_type: str, user_input: str):
    """Run the full smart pipeline and store everything in session_state."""

    progress = st.progress(0, text="Starting pipeline...")

    # ─── STEP 1: Extract / generate raw text ───
    progress.progress(5, text="🔄 Routing input...")
    raw_text = get_text_from_input(input_type, user_input, progress)
    if not raw_text or not raw_text.strip():
        progress.empty()
        return

    st.session_state["sp_raw_text"] = raw_text

    # ─── STEP 2: Detect language & topic ───
    progress.progress(35, text="🌍 Detecting language and topic...")
    st.session_state["sp_meta"] = detect_language_and_topic(raw_text)

    # ─── STEP 3: TL;DR + detailed summary ───
    progress.progress(50, text="📝 Generating TL;DR...")
    st.session_state["sp_tldr"] = summarize_text(raw_text, style="tldr")

    progress.progress(60, text="📝 Generating detailed summary...")
    if input_type == "🔍 Topic":
        # use topic-aware summarizer for topic input
        st.session_state["sp_summary"] = ai_summarize_topic(user_input, raw_text)
    else:
        st.session_state["sp_summary"] = summarize_text(raw_text, style="detailed")

    # ─── STEP 4: Keywords ───
    progress.progress(72, text="🔑 Extracting keywords...")
    st.session_state["sp_keywords"] = extract_keywords(raw_text, n=10)

    # ─── STEP 5: Paraphrase (top portion to save tokens) ───
    progress.progress(82, text="✍️ Paraphrasing...")
    st.session_state["sp_paraphrased"] = paraphrase_text(raw_text[:3500])

    # ─── STEP 6: Enhance ───
    progress.progress(92, text="🚀 AI enhancing...")
    st.session_state["sp_enhanced"] = enhance_text(raw_text[:3500])

    # ─── DONE ───
    progress.progress(100, text="✅ Pipeline complete!")
    progress.empty()


if run_btn:
    if not user_input or not user_input.strip():
        st.error("⚠️ Please provide an input.")
    elif not st.session_state.get("gemini_api_key", "").strip():
        st.error("⚠️ Please add your Gemini API key in the sidebar first.")
    else:
        # Clear previous state
        for k in ("sp_raw_text", "sp_meta", "sp_tldr", "sp_summary",
                  "sp_keywords", "sp_paraphrased", "sp_enhanced"):
            st.session_state.pop(k, None)
        try:
            run_pipeline(input_type, user_input.strip())
            if "sp_summary" in st.session_state:
                st.success("🎉 Smart Pipeline complete — scroll down to see all results!")
        except Exception as e:
            st.error(f"❌ Pipeline error: {e}")


# ─────────────────────────────────────────────
# PIPELINE STATUS VISUALIZATION
# ─────────────────────────────────────────────

if "sp_raw_text" in st.session_state:
    gradient_divider()
    st.markdown("### ✅ Pipeline Status")

    steps = [
        ("Input",         "sp_raw_text"),
        ("Meta",          "sp_meta"),
        ("TL;DR",         "sp_tldr"),
        ("Summary",       "sp_summary"),
        ("Keywords",      "sp_keywords"),
        ("Paraphrased",   "sp_paraphrased"),
        ("Enhanced",      "sp_enhanced"),
    ]

    cols = st.columns(len(steps))
    for col, (name, key) in zip(cols, steps):
        done = key in st.session_state and st.session_state.get(key)
        with col:
            color = "#10b981" if done else "#64748b"
            icon = "✓" if done else "○"
            st.markdown(
                f"""
                <div style="text-align:center; padding:0.8rem 0.4rem; border-radius:12px;
                            background:rgba(30,30,60,0.4); border:1px solid {color}40;">
                    <div style="font-size:1.4rem; color:{color}; font-weight:700;">{icon}</div>
                    <div style="font-size:0.8rem; color:#cbd5e1; margin-top:0.3rem;">{name}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )


# ─────────────────────────────────────────────
# RESULTS
# ─────────────────────────────────────────────

if "sp_meta" in st.session_state:
    meta = st.session_state["sp_meta"]
    raw_word_count = len(st.session_state.get("sp_raw_text", "").split())

    gradient_divider()
    st.markdown("### 📊 Overview")

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("🌍 Language",  meta.get("language", "Unknown"))
    m2.metric("🎯 Topic",     meta.get("topic", "Unknown"))
    m3.metric("📝 Words",     f"{raw_word_count:,}")
    m4.metric("🔑 Keywords",  len(st.session_state.get("sp_keywords", [])))


if "sp_tldr" in st.session_state and st.session_state["sp_tldr"]:
    st.markdown(
        f"""
        <div class="glass-card" style="border-left: 4px solid #fbbf24; margin-top:1rem;">
            <div style="color:#fbbf24; font-weight:700; font-size:0.85rem; text-transform:uppercase; letter-spacing:0.06em; margin-bottom:0.4rem;">
                💡 TL;DR
            </div>
            <div style="color:#e2e8f0; font-size:1.05rem; line-height:1.6;">
                {st.session_state["sp_tldr"]}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


if "sp_keywords" in st.session_state and st.session_state["sp_keywords"]:
    st.markdown("##### 🔑 Keywords")
    kw_html = " ".join(
        f'<span class="status-pill success" style="font-size:0.9rem; margin:4px;">{kw}</span>'
        for kw in st.session_state["sp_keywords"]
    )
    st.markdown(f'<div style="line-height:2.4;">{kw_html}</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────
# RESULTS — TABS
# ─────────────────────────────────────────────

if "sp_summary" in st.session_state:
    gradient_divider()
    st.markdown("### 📋 Results")

    tab_labels = ["📝 Summary"]
    if "sp_paraphrased" in st.session_state: tab_labels.append("✍️ Paraphrased")
    if "sp_enhanced"    in st.session_state: tab_labels.append("🚀 Enhanced")
    tab_labels.append("📄 Original")

    tabs = st.tabs(tab_labels)

    with tabs[0]:
        st.markdown(st.session_state["sp_summary"])

    idx = 1
    if "sp_paraphrased" in st.session_state:
        with tabs[idx]:
            st.markdown(
                f'<div class="content-box ai">{st.session_state["sp_paraphrased"]}</div>',
                unsafe_allow_html=True,
            )
        idx += 1

    if "sp_enhanced" in st.session_state:
        with tabs[idx]:
            st.markdown(
                f'<div class="content-box ai">{st.session_state["sp_enhanced"]}</div>',
                unsafe_allow_html=True,
            )
        idx += 1

    with tabs[idx]:
        raw = st.session_state.get("sp_raw_text", "")
        st.text(raw[:5000] + ("..." if len(raw) > 5000 else ""))


# ─────────────────────────────────────────────
# DOWNLOAD — COMPLETE REPORT
# ─────────────────────────────────────────────

def make_complete_report() -> bytes:
    """Build a single Word doc containing every pipeline output."""
    doc = Document()
    doc.add_heading("Smart Pipeline — Complete AI Report", level=1)

    meta = st.session_state.get("sp_meta", {})
    if meta:
        doc.add_paragraph(f"Language: {meta.get('language', 'Unknown')}")
        doc.add_paragraph(f"Topic: {meta.get('topic', 'Unknown')}")

    if st.session_state.get("sp_keywords"):
        doc.add_heading("Keywords", level=2)
        doc.add_paragraph(", ".join(st.session_state["sp_keywords"]))

    if st.session_state.get("sp_tldr"):
        doc.add_heading("TL;DR", level=2)
        doc.add_paragraph(st.session_state["sp_tldr"])

    if st.session_state.get("sp_summary"):
        doc.add_heading("Detailed Summary", level=2)
        for line in st.session_state["sp_summary"].split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("## "):
                doc.add_heading(line[3:], level=3)
            elif line.startswith(("- ", "* ")):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)

    if st.session_state.get("sp_paraphrased"):
        doc.add_heading("AI Paraphrased Version", level=2)
        doc.add_paragraph(st.session_state["sp_paraphrased"])

    if st.session_state.get("sp_enhanced"):
        doc.add_heading("AI Enhanced Version", level=2)
        doc.add_paragraph(st.session_state["sp_enhanced"])

    if st.session_state.get("sp_raw_text"):
        doc.add_heading("Original Source Text", level=2)
        raw = st.session_state["sp_raw_text"]
        for p in raw.split("\n\n"):
            if p.strip():
                doc.add_paragraph(p.strip())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


if "sp_summary" in st.session_state:
    gradient_divider()
    st.markdown("### 📥 Download Complete Report")
    st.download_button(
        "📦 Download Full AI Report (.docx)",
        data=make_complete_report(),
        file_name="smart_pipeline_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        type="primary",
    )

gradient_divider()
st.caption("Smart Pipeline · Full automation · Powered by Google Gemini 2.0 Flash")

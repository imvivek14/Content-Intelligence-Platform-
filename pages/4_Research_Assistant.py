"""
Research Assistant — Search arXiv papers and explain abstracts in plain English.
"""

import io
import xml.etree.ElementTree as ET
import requests
import streamlit as st
from docx import Document

from ai_helper import (
    explain_research_abstract,
    paraphrase_text,
    summarize_text,
    extract_keywords,
    render_api_key_sidebar,
)
from styles import inject_css, page_header, gradient_divider


# ─────────────────────────────────────────────
# PAGE
# ─────────────────────────────────────────────

st.set_page_config(page_title="Research Assistant", layout="wide", page_icon="📚")
inject_css()
render_api_key_sidebar()

page_header(
    icon="📚",
    title="AI Research Assistant",
    subtitle="Search live arXiv papers. Gemini explains complex abstracts in plain English and paraphrases them on demand.",
    badge="Module 04",
)


# ─────────────────────────────────────────────
# ARXIV API
# ─────────────────────────────────────────────

ATOM_NS = "{http://www.w3.org/2005/Atom}"


def fetch_arxiv(topic: str, max_results: int = 8) -> list[dict]:
    """Fetch papers from the arXiv API. Returns a list of paper dicts."""
    url = (
        f"http://export.arxiv.org/api/query"
        f"?search_query=all:{topic.replace(' ', '+')}"
        f"&start=0&max_results={max_results}"
        f"&sortBy=relevance&sortOrder=descending"
    )
    response = requests.get(url, timeout=15)
    response.raise_for_status()

    root = ET.fromstring(response.content)
    papers = []
    for entry in root.findall(f"{ATOM_NS}entry"):
        title_el     = entry.find(f"{ATOM_NS}title")
        summary_el   = entry.find(f"{ATOM_NS}summary")
        id_el        = entry.find(f"{ATOM_NS}id")
        published_el = entry.find(f"{ATOM_NS}published")

        title = (title_el.text.strip().replace("\n", " ") if title_el is not None else "Untitled")
        summary = (summary_el.text.strip().replace("\n", " ") if summary_el is not None else "No abstract.")
        link = id_el.text.strip() if id_el is not None else "#"
        published = published_el.text[:10] if published_el is not None else "Unknown"

        authors = []
        for author in entry.findall(f"{ATOM_NS}author"):
            name_el = author.find(f"{ATOM_NS}name")
            if name_el is not None and name_el.text:
                authors.append(name_el.text)

        papers.append({
            "title":     title,
            "authors":   ", ".join(authors[:4]) + (" et al." if len(authors) > 4 else ""),
            "summary":   summary,
            "link":      link,
            "published": published,
        })
    return papers


# ─────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────

topic = st.text_input(
    "🔬 Research topic",
    placeholder="e.g. neural networks, quantum computing, climate change, CRISPR",
    key="ra_topic",
)

c1, c2 = st.columns([1, 3])
with c1:
    max_papers = st.number_input("Papers", min_value=3, max_value=20, value=8, step=1)
with c2:
    st.write("")
    st.write("")
    fetch_btn = st.button("📡 Fetch Papers from arXiv", use_container_width=True, type="primary")

if fetch_btn:
    if not topic.strip():
        st.error("⚠️ Please enter a research topic.")
    else:
        try:
            with st.spinner(f"Searching arXiv for '{topic}'..."):
                papers = fetch_arxiv(topic.strip(), max_papers)
            if not papers:
                st.warning("No papers found. Try a different topic.")
            else:
                st.session_state["ra_papers"] = papers
                for k in ("ra_current_summary", "ra_current_title", "ra_explanation",
                          "ra_paraphrase", "ra_tldr", "ra_keywords"):
                    st.session_state.pop(k, None)
                st.success(f"✅ Found {len(papers)} papers.")
        except Exception as e:
            st.error(f"❌ Error fetching papers: {e}")


# ─────────────────────────────────────────────
# PAPER LIST
# ─────────────────────────────────────────────

if "ra_papers" in st.session_state:
    gradient_divider()
    st.markdown("### 📄 Papers")

    for i, paper in enumerate(st.session_state["ra_papers"]):
        with st.expander(f"📄 {i + 1}. {paper['title']}", expanded=(i == 0)):
            st.caption(f"**Authors:** {paper['authors']}  ·  **Published:** {paper['published']}")
            st.write(paper["summary"])
            ca, cb = st.columns([1, 1])
            with ca:
                st.markdown(f"[🔗 Read on arXiv]({paper['link']})")
            with cb:
                if st.button(f"🤖 Analyze this abstract", key=f"select_{i}", use_container_width=True):
                    st.session_state["ra_current_summary"] = paper["summary"]
                    st.session_state["ra_current_title"]   = paper["title"]
                    st.session_state["ra_current_link"]    = paper["link"]
                    for k in ("ra_explanation", "ra_paraphrase", "ra_tldr", "ra_keywords"):
                        st.session_state.pop(k, None)
                    st.rerun()


# ─────────────────────────────────────────────
# SELECTED PAPER ANALYSIS
# ─────────────────────────────────────────────

if "ra_current_summary" in st.session_state:
    gradient_divider()
    st.markdown(f"### 📌 Selected: {st.session_state['ra_current_title']}")
    st.markdown(
        f'<div class="content-box original">{st.session_state["ra_current_summary"]}</div>',
        unsafe_allow_html=True,
    )

    has_key = bool(st.session_state.get("gemini_api_key", "").strip())

    b1, b2, b3 = st.columns(3)
    explain_btn    = b1.button("🤖 Plain-English Explanation", use_container_width=True, disabled=not has_key)
    paraphrase_btn = b2.button("✍️ Paraphrase Abstract",       use_container_width=True, disabled=not has_key)
    full_btn       = b3.button("🎯 Full Analysis",              use_container_width=True, type="primary", disabled=not has_key)

    if not has_key:
        st.warning("⚠️ Add a Gemini API key in the sidebar to enable AI features.")

    if explain_btn:
        with st.spinner("🤖 Explaining..."):
            st.session_state["ra_explanation"] = explain_research_abstract(
                st.session_state["ra_current_title"],
                st.session_state["ra_current_summary"],
            )

    if paraphrase_btn:
        with st.spinner("✍️ Paraphrasing..."):
            st.session_state["ra_paraphrase"] = paraphrase_text(st.session_state["ra_current_summary"])

    if full_btn:
        progress = st.progress(0, text="Starting full analysis...")

        progress.progress(20, text="🤖 Generating plain-English explanation...")
        st.session_state["ra_explanation"] = explain_research_abstract(
            st.session_state["ra_current_title"],
            st.session_state["ra_current_summary"],
        )

        progress.progress(50, text="✍️ Paraphrasing abstract...")
        st.session_state["ra_paraphrase"] = paraphrase_text(st.session_state["ra_current_summary"])

        progress.progress(75, text="📝 Generating TL;DR...")
        st.session_state["ra_tldr"] = summarize_text(st.session_state["ra_current_summary"], style="tldr")

        progress.progress(95, text="🔑 Extracting keywords...")
        st.session_state["ra_keywords"] = extract_keywords(st.session_state["ra_current_summary"], n=8)

        progress.progress(100, text="✅ Done!")
        progress.empty()


# ─────────────────────────────────────────────
# RESULTS
# ─────────────────────────────────────────────

if "ra_tldr" in st.session_state:
    st.markdown(
        f"""
        <div class="glass-card" style="border-left: 4px solid #fbbf24;">
            <div style="color:#fbbf24; font-weight:700; font-size:0.85rem; text-transform:uppercase; letter-spacing:0.06em; margin-bottom:0.4rem;">
                💡 TL;DR
            </div>
            <div style="color:#e2e8f0; font-size:1.05rem; line-height:1.6;">
                {st.session_state["ra_tldr"]}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

if "ra_keywords" in st.session_state and st.session_state["ra_keywords"]:
    st.markdown("##### 🔑 Key Concepts")
    kw_html = " ".join(
        f'<span class="status-pill success" style="margin:4px;">{kw}</span>'
        for kw in st.session_state["ra_keywords"]
    )
    st.markdown(f'<div style="line-height:2.4;">{kw_html}</div>', unsafe_allow_html=True)

if "ra_explanation" in st.session_state:
    gradient_divider()
    st.markdown("### 🤖 Plain-English Explanation")
    st.markdown(st.session_state["ra_explanation"])

if "ra_paraphrase" in st.session_state:
    gradient_divider()
    st.markdown("### ✍️ AI Paraphrased Abstract")
    st.markdown(
        f'<div class="content-box ai">{st.session_state["ra_paraphrase"]}</div>',
        unsafe_allow_html=True,
    )


# ─────────────────────────────────────────────
# DOWNLOAD
# ─────────────────────────────────────────────

def make_doc_bytes(text: str, title: str = "") -> bytes:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for p in (text or "").split("\n"):
        if p.strip():
            doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


if "ra_current_summary" in st.session_state:
    gradient_divider()
    st.markdown("### 📥 Download")
    title = st.session_state.get("ra_current_title", "Paper")
    cols = st.columns(3)

    with cols[0]:
        st.download_button(
            "📄 Original Abstract",
            make_doc_bytes(st.session_state["ra_current_summary"], title),
            file_name="abstract_original.docx",
            use_container_width=True,
        )
    if "ra_explanation" in st.session_state:
        with cols[1]:
            st.download_button(
                "🤖 AI Explanation",
                make_doc_bytes(st.session_state["ra_explanation"], f"AI Explanation: {title}"),
                file_name="ai_explanation.docx",
                use_container_width=True,
            )
    if "ra_paraphrase" in st.session_state:
        with cols[2]:
            st.download_button(
                "✍️ AI Paraphrase",
                make_doc_bytes(st.session_state["ra_paraphrase"], f"Paraphrased: {title}"),
                file_name="ai_paraphrase.docx",
                use_container_width=True,
            )

gradient_divider()
st.caption("Research Assistant · arXiv API · Explained by Google Gemini 2.0 Flash")

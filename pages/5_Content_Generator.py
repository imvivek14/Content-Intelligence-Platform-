"""
Content Generator — Generate fresh structured content from any prompt.

Fixes in this revision:
  • Explicit `index=0` on dropdowns so default values render visibly.
  • Errors from AI calls are shown in a red error box (not hidden).
  • Pre-flight key check so the user gets clear feedback before submitting.
"""

import io
import streamlit as st
from docx import Document

from ai_helper import (
    generate_content,
    paraphrase_text,
    summarize_text,
    extract_keywords,
    enhance_text,
    render_api_key_sidebar,
    _resolve_api_key,
)
from styles import inject_css, page_header, gradient_divider, show_ai_error


# ─────────────────────────────────────────────
# PAGE
# ─────────────────────────────────────────────

st.set_page_config(page_title="AI Content Generator", layout="wide", page_icon="✍️")
inject_css()
render_api_key_sidebar()

page_header(
    icon="✍️",
    title="AI Content Generator",
    subtitle="From a single prompt: stories, essays, blog posts, explanations — with adjustable tone, length, and audience.",
    badge="Module 05",
)


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def is_error(s: str) -> bool:
    return bool(s) and s.lstrip().startswith("⚠️ AI request failed")


# ─────────────────────────────────────────────
# QUICK-START EXAMPLES
# ─────────────────────────────────────────────

st.markdown("##### 💡 Try one of these prompts")

examples = [
    "Write a short story about a robot learning to feel emotions",
    "Explain quantum computing in simple terms",
    "Advantages and disadvantages of artificial intelligence",
    "Write an essay about climate change and its future impact",
    "Compose a motivational LinkedIn post about career growth",
    "Describe the journey from idea to startup launch",
]

ex_cols = st.columns(3)
for i, ex in enumerate(examples):
    with ex_cols[i % 3]:
        if st.button(ex, use_container_width=True, key=f"ex_{i}"):
            st.session_state["cg_prompt_value"] = ex
            st.rerun()

gradient_divider()


# ─────────────────────────────────────────────
# INPUT
# ─────────────────────────────────────────────

prompt = st.text_area(
    "💬 Your prompt",
    value=st.session_state.get("cg_prompt_value", ""),
    height=120,
    placeholder='e.g. "Write a story about AI", "Explain blockchain", "Advantages of remote work"',
    key="cg_prompt_input",
)

# Explicit index=0 so a value is always selected and visible
TONES     = ["Professional", "Creative", "Simple & Clear", "Academic", "Casual", "Persuasive", "Humorous"]
LENGTHS   = ["Short (1–2 paragraphs)", "Medium (3–5 paragraphs)", "Long (detailed article)"]
AUDIENCES = ["General", "Beginners", "Experts", "Students", "Business", "Children"]

c1, c2, c3 = st.columns(3)
tone     = c1.selectbox("🎨 Tone",     TONES,     index=0, key="cg_tone")
length   = c2.selectbox("📏 Length",   LENGTHS,   index=1, key="cg_length")  # default to Medium
audience = c3.selectbox("👥 Audience", AUDIENCES, index=0, key="cg_audience")


b1, b2, b3 = st.columns(3)
generate_btn      = b1.button("⚡ Generate",                use_container_width=True, type="primary", disabled=not prompt.strip())
paraphrase_btn    = b2.button("✍️ Paraphrase Output",       use_container_width=True, disabled="cg_content" not in st.session_state)
auto_pipeline_btn = b3.button("🎯 Generate + Auto-Analyze", use_container_width=True, disabled=not prompt.strip())


has_key = bool(_resolve_api_key())


def needs_key() -> bool:
    if not has_key:
        st.error("⚠️ No Gemini API key. Paste your key in the sidebar and click **🔌 Test Connection** to verify.")
        return True
    return False


def build_full_prompt(user_prompt: str) -> str:
    return (
        f"Tone: {tone}\n"
        f"Length: {length}\n"
        f"Target audience: {audience}\n\n"
        f"Request: {user_prompt}"
    )


# ─────────────────────────────────────────────
# GENERATE
# ─────────────────────────────────────────────

if generate_btn:
    if needs_key():
        pass
    else:
        with st.spinner("🤖 Gemini is generating content..."):
            content = generate_content(build_full_prompt(prompt))
            st.session_state["cg_content"] = content
            for k in ("cg_paraphrased", "cg_summary", "cg_keywords", "cg_enhanced"):
                st.session_state.pop(k, None)
        if is_error(content):
            st.error("Generation failed — see the error message below.")
        else:
            st.success("✅ Content generated!")


if paraphrase_btn and "cg_content" in st.session_state and not needs_key():
    with st.spinner("✍️ AI paraphrasing..."):
        st.session_state["cg_paraphrased"] = paraphrase_text(st.session_state["cg_content"])


if auto_pipeline_btn and not needs_key():
    progress = st.progress(0, text="Starting...")

    progress.progress(20, text="🤖 Generating content...")
    st.session_state["cg_content"] = generate_content(build_full_prompt(prompt))

    if is_error(st.session_state["cg_content"]):
        progress.empty()
        st.error("Generation failed — see error below.")
    else:
        progress.progress(45, text="🚀 AI enhancing...")
        st.session_state["cg_enhanced"] = enhance_text(st.session_state["cg_content"])

        progress.progress(70, text="📝 Generating summary...")
        st.session_state["cg_summary"] = summarize_text(st.session_state["cg_content"], style="concise")

        progress.progress(90, text="🔑 Extracting keywords...")
        st.session_state["cg_keywords"] = extract_keywords(st.session_state["cg_content"], n=8)

        progress.progress(100, text="✅ Done!")
        progress.empty()
        st.success("🎉 Content + analysis ready!")


# ─────────────────────────────────────────────
# RESULTS
# ─────────────────────────────────────────────

if "cg_content" in st.session_state:
    gradient_divider()

    content = st.session_state["cg_content"]

    # If the generation itself failed — show the error prominently and stop
    if is_error(content):
        show_ai_error(content)
    else:
        # Quick metrics chips
        word_count = len(content.split())
        char_count = len(content)
        st.markdown(
            f"""
            <div style="margin-bottom: 0.8rem;">
                <span class="status-pill success">📝 {word_count} words</span>
                <span class="status-pill success">🔠 {char_count} characters</span>
                <span class="status-pill success">🎨 {tone}</span>
                <span class="status-pill success">👥 {audience}</span>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Summary card
        if "cg_summary" in st.session_state:
            summary = st.session_state["cg_summary"]
            if is_error(summary):
                show_ai_error(summary)
            else:
                st.markdown(
                    f"""
                    <div class="glass-card" style="border-left: 4px solid #22d3ee;">
                        <div style="color:#22d3ee; font-weight:700; font-size:0.85rem; text-transform:uppercase; letter-spacing:0.06em; margin-bottom:0.4rem;">
                            📝 AI Summary
                        </div>
                        <div style="color:#e2e8f0; line-height:1.6;">{summary}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

        # Keywords
        if "cg_keywords" in st.session_state and st.session_state["cg_keywords"]:
            kw_html = " ".join(
                f'<span class="status-pill success" style="margin:4px;">{kw}</span>'
                for kw in st.session_state["cg_keywords"]
            )
            st.markdown(f'<div style="line-height:2.4; margin-bottom:1rem;">{kw_html}</div>', unsafe_allow_html=True)

        # Tabs for views
        tab_labels = ["📄 Generated"]
        if "cg_enhanced" in st.session_state and not is_error(st.session_state["cg_enhanced"]):
            tab_labels.append("🚀 AI Enhanced")
        if "cg_paraphrased" in st.session_state and not is_error(st.session_state["cg_paraphrased"]):
            tab_labels.append("✍️ Paraphrased")

        tabs = st.tabs(tab_labels)
        with tabs[0]:
            st.markdown(content)

        idx = 1
        if "cg_enhanced" in st.session_state and not is_error(st.session_state["cg_enhanced"]):
            with tabs[idx]:
                st.markdown(st.session_state["cg_enhanced"])
            idx += 1
        if "cg_paraphrased" in st.session_state and not is_error(st.session_state["cg_paraphrased"]):
            with tabs[idx]:
                st.markdown(st.session_state["cg_paraphrased"])


# ─────────────────────────────────────────────
# DOWNLOAD
# ─────────────────────────────────────────────

def make_docx(text: str, title: str = "Generated Content") -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in (text or "").split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


if "cg_content" in st.session_state and not is_error(st.session_state["cg_content"]):
    gradient_divider()
    st.markdown("### 📥 Download")

    cols = st.columns(3)
    with cols[0]:
        st.download_button(
            "📄 Generated",
            data=make_docx(st.session_state["cg_content"], "AI Generated Content"),
            file_name="ai_generated.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    if "cg_enhanced" in st.session_state and not is_error(st.session_state["cg_enhanced"]):
        with cols[1]:
            st.download_button(
                "🚀 Enhanced",
                data=make_docx(st.session_state["cg_enhanced"], "AI Enhanced Content"),
                file_name="ai_enhanced.docx",
                use_container_width=True,
            )
    if "cg_paraphrased" in st.session_state and not is_error(st.session_state["cg_paraphrased"]):
        with cols[2]:
            st.download_button(
                "✍️ Paraphrased",
                data=make_docx(st.session_state["cg_paraphrased"], "AI Paraphrased Content"),
                file_name="ai_paraphrased.docx",
                use_container_width=True,
            )

gradient_divider()
st.caption("Content Generator · Powered by Google Gemini")

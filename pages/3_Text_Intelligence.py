"""
Text Intelligence — Grammar fix, readability, sentiment, AI enhancement.

Key change: when the API key is connected, AI is the PRIMARY action
(not the rule-based "Quick Fix"). Real errors are shown, never hidden.
"""

from __future__ import annotations
import io
import re
from collections import Counter
import streamlit as st
from docx import Document

from ai_helper import (
    enhance_text,
    grammar_check_detailed,
    extract_keywords,
    summarize_text,
    paraphrase_text,
    render_api_key_sidebar,
    _resolve_api_key,
)
from styles import inject_css, page_header, gradient_divider, show_ai_error


st.set_page_config(page_title="Text Intelligence", layout="wide", page_icon="🧠")
inject_css()
render_api_key_sidebar()

page_header(
    icon="🧠",
    title="Text Intelligence",
    subtitle="AI-powered grammar correction, readability metrics, sentiment, and full text enhancement.",
    badge="Module 03",
)


# ─────────────────────────────────────────────
# STATIC GRAMMAR ENGINE (expanded — used as fallback)
# ─────────────────────────────────────────────

SPELLING_MAP = {
    # The screenshot's misspelling — fixed
    "yasterday": "yesterday", "yestarday": "yesterday", "yeasterday": "yesterday",
    "tommorow": "tomorrow", "tommorrow": "tomorrow",
    # Common misspellings
    "recieve": "receive", "recieved": "received", "recieves": "receives",
    "occured": "occurred", "occuring": "occurring", "occurance": "occurrence",
    "seperate": "separate", "seperated": "separated", "seperately": "separately",
    "definately": "definitely", "definatly": "definitely",
    "recomend": "recommend", "reccomend": "recommend",
    "accomodate": "accommodate", "untill": "until",
    "becuase": "because", "becouse": "because", "bcoz": "because",
    "teh": "the", "thier": "their", "wierd": "weird",
    "neccessary": "necessary", "begining": "beginning", "beggining": "beginning",
    "occassion": "occasion", "independant": "independent",
    "truely": "truly", "existance": "existence",
    "relevent": "relevant", "succesful": "successful",
    "wether": "whether", "wich": "which", "writting": "writing",
    "alot": "a lot", "everytime": "every time",
    "noone": "no one", "infact": "in fact", "incase": "in case",
    # Past tense errors
    "buyed": "bought", "goed": "went", "doed": "did",
    "runned": "ran", "writed": "wrote", "speaked": "spoke",
    "tooked": "took", "sayed": "said", "comed": "came",
    # Other common slips
    "wanna": "want to", "gonna": "going to", "gotta": "got to",
    "kinda": "kind of", "sorta": "sort of", "outta": "out of",
}


GRAMMAR_RULES: list[tuple[str, str]] = [
    (r'\b(\w+)\s+\1\b', r'\1'),                # repeated word
    (r'\bi\b', 'I'),                           # lowercase i → I
    (r'([.!?,;])([A-Za-z])', r'\1 \2'),        # space after punctuation
    (r' {2,}', ' '),                           # collapse multiple spaces
    (r"\bdont\b", "don't"), (r"\bcant\b", "can't"), (r"\bwont\b", "won't"),
    (r"\bwouldnt\b", "wouldn't"), (r"\bcouldnt\b", "couldn't"),
    (r"\bshouldnt\b", "shouldn't"), (r"\bisnt\b", "isn't"),
    (r"\barent\b", "aren't"), (r"\bwasnt\b", "wasn't"),
    (r"\bwerent\b", "weren't"), (r"\bhavent\b", "haven't"),
    (r"\bhasnt\b", "hasn't"), (r"\bhadnt\b", "hadn't"),
    (r"\bim\b", "I'm"), (r"\bive\b", "I've"),
    # 'a' → 'an' before vowel sounds
    (r'\ba ([aeiouAEIOU])', r'an \1'),
    (r'\ban ([^aeiouAEIOU\s])', r'a \1'),
]


def fix_spelling(text: str) -> str:
    def repl(match: re.Match) -> str:
        token = match.group(0)
        replacement = SPELLING_MAP.get(token.lower())
        if not replacement:
            return token
        if token[0].isupper():
            replacement = replacement[0].upper() + replacement[1:]
        return replacement
    return re.sub(r"\b[A-Za-z]+\b", repl, text)


def fix_grammar_rules(text: str) -> str:
    for pattern, replacement in GRAMMAR_RULES:
        text = re.sub(pattern, replacement, text)
    return text


def fix_capitalization(text: str) -> str:
    sentences = re.split(r'(?<=[.!?])\s+', text)
    fixed = []
    for s in sentences:
        s = s.strip()
        if s:
            fixed.append(s[0].upper() + s[1:])
    return " ".join(fixed)


def fix_punctuation(text: str) -> str:
    text = text.strip()
    if text and text[-1] not in ".!?":
        text += "."
    return re.sub(r'\s+([.,!?;:])', r'\1', text)


def remove_repeated_words(text: str) -> str:
    return re.sub(r'\b(\w+)( \1\b)+', r'\1', text, flags=re.IGNORECASE)


def quick_grammar_fix(text: str) -> str:
    text = fix_spelling(text)
    text = remove_repeated_words(text)
    text = fix_grammar_rules(text)
    text = fix_capitalization(text)
    text = fix_punctuation(text)
    return text


# ─────────────────────────────────────────────
# READABILITY & SENTIMENT (pure-python, no NLTK)
# ─────────────────────────────────────────────

POSITIVE_WORDS = {
    "good", "great", "excellent", "amazing", "wonderful", "fantastic", "love", "best",
    "happy", "joy", "perfect", "beautiful", "awesome", "brilliant", "outstanding",
    "superb", "delightful", "marvellous", "pleasant", "positive", "success",
}
NEGATIVE_WORDS = {
    "bad", "terrible", "awful", "horrible", "worst", "hate", "ugly", "poor",
    "sad", "angry", "disappointing", "fail", "wrong", "broken", "useless",
    "boring", "annoying", "negative", "problem", "difficult", "issue",
}


def count_syllables(word: str) -> int:
    word = word.lower()
    count = len(re.findall(r'[aeiou]+', word))
    if word.endswith("e") and count > 1:
        count -= 1
    return max(1, count)


def simple_sentiment(text: str) -> tuple[str, float]:
    words = re.findall(r'[a-zA-Z]+', text.lower())
    if not words:
        return ("😐 Neutral", 0.0)
    pos = sum(1 for w in words if w in POSITIVE_WORDS)
    neg = sum(1 for w in words if w in NEGATIVE_WORDS)
    score = (pos - neg) / max(len(words), 1)
    if score > 0.02:
        return ("😊 Positive", score)
    if score < -0.02:
        return ("😟 Negative", score)
    return ("😐 Neutral", score)


def compute_metrics(text: str) -> dict:
    words = text.split()
    sentences = [s for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]

    word_count = len([w for w in words if any(c.isalpha() for c in w)])
    sentence_count = max(len(sentences), 1)
    avg_words = round(word_count / sentence_count, 2)

    syllables = sum(count_syllables(w) for w in words if any(c.isalpha() for c in w))
    if word_count > 0 and sentence_count > 0:
        flesch = 206.835 - 1.015 * (word_count / sentence_count) - 84.6 * (syllables / max(word_count, 1))
        flesch = max(0, min(100, round(flesch, 1)))
    else:
        flesch = 0

    sentiment_label, _ = simple_sentiment(text)

    stopwords = {
        "the", "a", "an", "is", "are", "was", "were", "and", "or", "but", "in",
        "on", "at", "to", "for", "of", "with", "it", "this", "that", "be", "by",
        "as", "from", "have", "has", "had", "i", "you", "he", "she", "we", "they",
        "my", "your", "his", "her", "our", "their", "not", "no", "do", "does",
    }
    freq = Counter(
        w.lower() for w in words
        if any(c.isalpha() for c in w) and w.lower() not in stopwords and len(w) > 2
    )

    return {
        "word_count": word_count,
        "sentence_count": sentence_count,
        "avg_words_per_sentence": avg_words,
        "flesch_score": flesch,
        "sentiment": sentiment_label,
        "top_words": freq.most_common(8),
    }


def readability_label(score: float) -> str:
    if score >= 80: return "Very Easy"
    if score >= 70: return "Easy"
    if score >= 60: return "Standard"
    if score >= 50: return "Fairly Difficult"
    if score >= 30: return "Difficult"
    return "Very Difficult"


# ─────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────

text = st.text_area(
    "📝 Paste your text",
    height=180,
    placeholder="e.g. i went to the the store yasterday and buyed some apple. it help me alot.",
    key="ti_input",
)

has_key = bool(_resolve_api_key())

# Action buttons — order PRIORITIZES AI when key is set
if has_key:
    c1, c2, c3, c4 = st.columns(4)
    full_btn       = c1.button("🎯 Full AI Analysis", use_container_width=True, type="primary", disabled=not text.strip())
    enhance_btn    = c2.button("🚀 AI Enhance",       use_container_width=True, disabled=not text.strip())
    ai_grammar_btn = c3.button("🤖 AI Grammar Audit", use_container_width=True, disabled=not text.strip())
    analyze_btn    = c4.button("⚡ Quick Fix (rules)", use_container_width=True, disabled=not text.strip())
else:
    st.info("🔑 Add your Gemini API key in the sidebar (and click **Test Connection**) to unlock AI features. Quick Fix works without a key.")
    c1 = st.columns(1)[0]
    analyze_btn    = c1.button("⚡ Quick Fix (rules only)", use_container_width=True, disabled=not text.strip(), type="primary")
    full_btn = enhance_btn = ai_grammar_btn = False


# ─────────────────────────────────────────────
# ACTIONS
# ─────────────────────────────────────────────

def is_error(s: str) -> bool:
    """Detect if a returned string is an AI error rather than real output."""
    return bool(s) and s.lstrip().startswith("⚠️ AI request failed")


if analyze_btn and text.strip():
    with st.spinner("Analyzing..."):
        st.session_state["ti_original"]  = text
        st.session_state["ti_corrected"] = quick_grammar_fix(text)
        st.session_state["ti_metrics"]   = compute_metrics(text)
        for k in ("ti_enhanced", "ti_grammar_report", "ti_summary", "ti_keywords"):
            st.session_state.pop(k, None)


if ai_grammar_btn and text.strip():
    with st.spinner("🤖 AI auditing grammar..."):
        st.session_state["ti_original"] = text
        st.session_state["ti_grammar_report"] = grammar_check_detailed(text)
        st.session_state["ti_metrics"] = compute_metrics(text)


if enhance_btn and text.strip():
    with st.spinner("🚀 AI enhancing..."):
        st.session_state["ti_original"] = text
        st.session_state["ti_enhanced"] = enhance_text(text)
        st.session_state["ti_metrics"] = compute_metrics(text)


if full_btn and text.strip():
    progress = st.progress(0, text="Starting full AI analysis...")
    st.session_state["ti_original"] = text

    progress.progress(15, text="Computing metrics...")
    st.session_state["ti_metrics"] = compute_metrics(text)

    progress.progress(28, text="⚡ Quick rule-based fix...")
    st.session_state["ti_corrected"] = quick_grammar_fix(text)

    progress.progress(45, text="🤖 AI grammar audit...")
    st.session_state["ti_grammar_report"] = grammar_check_detailed(text)

    progress.progress(65, text="🚀 AI enhancement...")
    st.session_state["ti_enhanced"] = enhance_text(text)

    progress.progress(82, text="🔑 Keyword extraction...")
    st.session_state["ti_keywords"] = extract_keywords(text, n=8)

    progress.progress(94, text="📝 AI summary...")
    st.session_state["ti_summary"] = summarize_text(text, style="concise")

    progress.progress(100, text="✅ Done!")
    progress.empty()
    st.success("🎉 Full AI analysis complete!")


# ─────────────────────────────────────────────
# RESULTS — METRICS
# ─────────────────────────────────────────────

if "ti_metrics" in st.session_state:
    m = st.session_state["ti_metrics"]
    gradient_divider()
    st.markdown("### 📊 Text Metrics")

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("📝 Words", m["word_count"])
    c2.metric("📖 Sentences", m["sentence_count"])
    c3.metric("📏 Avg Words/Sent", m["avg_words_per_sentence"])
    c4.metric("📊 Flesch", m["flesch_score"], readability_label(m["flesch_score"]))
    c5.metric("💬 Sentiment", m["sentiment"])

    if m["top_words"]:
        st.markdown("##### 🔝 Most-used words")
        chips_html = " ".join(
            f'<span class="status-pill success" style="margin:4px;">{w} <span style="opacity:0.7;">({c})</span></span>'
            for w, c in m["top_words"]
        )
        st.markdown(f'<div style="line-height:2.4;">{chips_html}</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────
# RESULTS — AI KEYWORDS
# ─────────────────────────────────────────────

if "ti_keywords" in st.session_state and st.session_state["ti_keywords"]:
    st.markdown("##### 🔑 AI-Extracted Keywords")
    kw_html = " ".join(
        f'<span class="status-pill success" style="margin:4px;">{kw}</span>'
        for kw in st.session_state["ti_keywords"]
    )
    st.markdown(f'<div style="line-height:2.4;">{kw_html}</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────
# RESULTS — AI SUMMARY (with error visibility)
# ─────────────────────────────────────────────

if "ti_summary" in st.session_state and st.session_state["ti_summary"]:
    gradient_divider()
    st.markdown("### 📝 AI Summary")
    summary = st.session_state["ti_summary"]
    if is_error(summary):
        show_ai_error(summary)
    else:
        st.markdown(f'<div class="content-box ai">{summary}</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────
# RESULTS — QUICK GRAMMAR FIX
# ─────────────────────────────────────────────

if "ti_corrected" in st.session_state:
    gradient_divider()
    st.markdown("### ⚡ Quick Grammar Fix (rule-based)")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("#### 📄 Original")
        st.markdown(f'<div class="content-box original">{st.session_state["ti_original"]}</div>', unsafe_allow_html=True)
    with col2:
        st.markdown("#### ✅ Fixed")
        st.markdown(f'<div class="content-box ai">{st.session_state["ti_corrected"]}</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────
# RESULTS — AI GRAMMAR AUDIT
# ─────────────────────────────────────────────

if "ti_grammar_report" in st.session_state:
    gradient_divider()
    st.markdown("### 🤖 AI Grammar Audit")
    report = st.session_state["ti_grammar_report"]
    if is_error(report):
        show_ai_error(report)
    else:
        st.markdown(report)


# ─────────────────────────────────────────────
# RESULTS — AI ENHANCEMENT
# ─────────────────────────────────────────────

if "ti_enhanced" in st.session_state:
    gradient_divider()
    st.markdown("### 🚀 AI-Enhanced Text")
    enhanced = st.session_state["ti_enhanced"]
    if is_error(enhanced):
        show_ai_error(enhanced)
    else:
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("#### Original")
            st.markdown(f'<div class="content-box original">{st.session_state["ti_original"]}</div>', unsafe_allow_html=True)
        with col2:
            st.markdown("#### AI Enhanced")
            st.markdown(f'<div class="content-box ai">{enhanced}</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────
# DOWNLOAD
# ─────────────────────────────────────────────

def make_doc_bytes(t: str, title: str = ""):
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for p in (t or "").split("\n"):
        if p.strip():
            doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


if any(k in st.session_state for k in ("ti_original", "ti_corrected", "ti_enhanced")):
    gradient_divider()
    st.markdown("### 📥 Download")
    cols = st.columns(3)
    if "ti_original" in st.session_state:
        with cols[0]:
            st.download_button(
                "📄 Original",
                make_doc_bytes(st.session_state["ti_original"], "Original Text"),
                file_name="original.docx",
                use_container_width=True,
            )
    if "ti_corrected" in st.session_state:
        with cols[1]:
            st.download_button(
                "✅ Quick Fixed",
                make_doc_bytes(st.session_state["ti_corrected"], "Grammar Corrected"),
                file_name="corrected.docx",
                use_container_width=True,
            )
    if "ti_enhanced" in st.session_state and not is_error(st.session_state["ti_enhanced"]):
        with cols[2]:
            st.download_button(
                "🚀 AI Enhanced",
                make_doc_bytes(st.session_state["ti_enhanced"], "AI Enhanced Text"),
                file_name="ai_enhanced.docx",
                use_container_width=True,
            )

gradient_divider()
st.caption("Text Intelligence · Pure-Python rule engine + Google Gemini")
